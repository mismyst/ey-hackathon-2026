#!/usr/bin/env python3
"""Generate SHELDRA enterprise architecture document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Global Styles ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

code_style = doc.styles.add_style('CodeBlock', 1)  # paragraph
code_style.font.name = 'Consolas'
code_style.font.size = Pt(8.5)
code_style.paragraph_format.space_before = Pt(6)
code_style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    h.font.name = 'Calibri'

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.style.font.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_para(text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    return p

def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

def add_code(code_text):
    lines = code_text.strip().split('\n')
    for line in lines:
        p = doc.add_paragraph(style='CodeBlock')
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(8.5)

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p

def add_subsection_heading(text):
    """Add a subsection heading (bold, no number)."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x33, 0x55, 0x88)
    return p

# ── Title Page ────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('SHELDRA')
r.bold = True
r.font.size = Pt(42)
r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Safety Holographic Enhanced Learning & Decision-Reality Assistant')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Enterprise Architecture Document')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x6B, 0x8F, 0xC6)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'ET AI Hackathon 2026 \u2014 Industrial Safety Intelligence Challenge\n')
meta.add_run(f'Version 1.0 \u2014 {datetime.date.today().strftime("%B %d, %Y")}\n')
meta.add_run(f'Classification: Public \u2014 Architecture Overview')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (Manual)
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
sections = [
    '1. Executive Summary', '2. Problem Statement', '3. Vision',
    '4. Functional Requirements', '5. Non-Functional Requirements',
    '6. Overall System Architecture', '7. SHELDRA Intelligence Engine',
    '8. Decision Tree Learning System', '9. Computer Vision Pipeline',
    '10. Time Series Analytics', '11. Knowledge Graph',
    '12. Retrieval Augmented Generation', '13. Data Flow Architecture',
    '14. Database Design', '15. API Architecture', '16. Frontend Architecture',
    '17. Security', '18. Deployment Architecture', '19. Technology Stack',
    '20. AI Models', '21. Performance', '22. Scalability',
    '23. Demo Flow', '24. Roadmap', '25. Risks', '26. Future Enhancements'
]
for s in sections:
    p = doc.add_paragraph(s)
    p.paragraph_format.space_after = Pt(1)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════
add_heading('1. Executive Summary', level=1)

add_para(
    'SHELDRA is a unified AI system for industrial safety intelligence. Unlike conventional '
    'approaches that layer AI features onto dashboards, SHELDRA is a single Intelligence Engine '
    'with 11 internal modules that together perceive, reason, learn, and respond to industrial '
    'safety events in real time.'
)
add_para(
    'The system ingests heterogeneous data streams \u2014 camera frames, IoT sensor telemetry, '
    'worker interaction history, safety documentation \u2014 through a single Event Orchestrator '
    'that routes each input to the appropriate internal module. An LLM Reasoning Core generates '
    'personalized responses, a Decision Tree Engine maintains self-improving hazard resolution '
    'trees, and a Memory Manager preserves both session state and long-term interaction history. '
    'Every decision is logged by the Explainability Engine and validated by Safety Guardrails '
    'before delivery.'
)
add_para(
    'The architecture follows a strict layering: the SHELDRA Intelligence Engine has zero '
    'knowledge of the frontend. All interactions occur through a FastAPI API layer. The '
    'presentation layer (dashboard, 3D avatar, voice, VR) is interchangeable without modifying '
    'the engine. This separation enables enterprise deployment scenarios from single-facility '
    'on-premise to global multi-region SaaS.'
)
add_para(
    'Key technical decisions include: Apache Kafka for durable event streaming, Neo4j for '
    'relationship-rich knowledge representation, Qdrant for hybrid vector search, Llama 3 8B '
    'as the reasoning core, and ONNX Runtime for cross-platform model serving. All components '
    'are containerized for deployment consistency.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════════════════
add_heading('2. Problem Statement', level=1)

add_para(
    'Industrial safety monitoring today suffers from four architectural failures that no '
    'single-vendor solution adequately addresses.'
)

add_subsection_heading('2.1 Latent Detection')
add_para(
    'Current systems detect incidents after they occur. Data is collected, stored, and analyzed '
    'retrospectively. By the time a dashboard indicator changes, the event has already '
    'happened. This is a fundamental architectural limitation of batch-oriented data pipelines '
    'that cannot support real-time intervention.'
)

add_subsection_heading('2.2 Modality Isolation')
add_para(
    'Computer vision, IoT telemetry, and safety documentation are managed by separate systems '
    'with no shared context. A worker who removes PPE while a nearby machine shows anomalous '
    'vibration and gas readings represents a compound risk that no isolated system can detect. '
    'The absence of a unified intelligence layer means compound risks are invisible.'
)

add_subsection_heading('2.3 Static Knowledge')
add_para(
    'Safety checklists and standard operating procedures are static documents. When an incident '
    'reveals a gap in a procedure, updating that procedure for all workers is an organizational '
    'process that takes weeks or months. There is no feedback loop between incident outcome and '
    'procedure improvement.'
)

add_subsection_heading('2.4 One-Size-Fits-All Guidance')
add_para(
    'Safety instructions are delivered identically to every worker regardless of experience, '
    'language, learning speed, or current emotional state. A 20-year veteran and a first-day '
    'intern receive the same procedure document. This increases cognitive load for experienced '
    'workers and provides insufficient guidance for novices.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 3. VISION
# ═══════════════════════════════════════════════════════════════════════════
add_heading('3. Vision', level=1)
add_para(
    'A single AI system that perceives every signal in an industrial environment, reasons over '
    'fused data with a Large Language Model, maintains a living knowledge graph of entities '
    'and relationships, generates self-improving hazard resolution trees, and delivers '
    'personalized, explainable safety guidance through any interface \u2014 from a 3D '
    'holographic avatar to a simple text alert.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 4. FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════════
add_heading('4. Functional Requirements', level=1)

add_para('The platform shall satisfy the following functional requirements:', bold=True)

func_reqs = [
    ['FR1', 'Real-Time Coaching', 'SHELDRA shall provide personalized safety guidance to workers within 3 seconds of any triggering event, adapting its response to the worker\u2019s profile, language, and current emotional state.'],
    ['FR2', 'Hazard Resolution Trees', 'The system shall generate, maintain, and automatically update decision trees for hazard resolution. Trees shall version on every update with the ability to roll back to any prior version.'],
    ['FR3', 'Multi-Modal Perception', 'SHELDRA shall ingest and fuse data from camera feeds (object detection, pose estimation), IoT sensors (temperature, vibration, gas, noise), and safety documentation (procedures, regulations, incident reports).'],
    ['FR4', 'Personalization', 'The system shall maintain individual worker profiles that include experience level, learning speed, certifications, interaction history, and emotional state. SHELDRA\u2019s responses shall adapt along six dimensions: content depth, vocabulary, speech rate, tone, feedback style, and language.'],
    ['FR5', 'Explainability', 'Every SHELDRA decision shall be traceable to specific observations, rules, and reasoning steps. The system shall produce natural language explanations, confidence scores with provenance, and structured decision traces.'],
    ['FR6', 'Human-in-the-Loop Oversight', 'Supervisors shall be able to monitor all SHELDRA interactions, override recommendations, and escalate decisions. Every HITL action shall be logged for audit.'],
    ['FR7', 'Knowledge Grounding', 'All SHELDRA guidance shall be grounded in retrievable safety documentation. Responses shall include citations to specific document sections.'],
    ['FR8', 'Avatar Interface', 'SHELDRA shall render as a 3D holographic avatar with lip-sync, gesture animation, and emotional expression. The avatar shall be capable of appearing in web dashboard, mobile, and VR environments.'],
]
add_table(['ID', 'Requirement', 'Description'], func_reqs)

# ═══════════════════════════════════════════════════════════════════════════
# 5. NON-FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════════
add_heading('5. Non-Functional Requirements', level=1)

nfr = [
    ['NFR1', 'Latency', 'End-to-end response time < 3 seconds for 95th percentile under normal load. SHELDRA streaming response first token < 500ms.'],
    ['NFR2', 'Throughput', 'System shall sustain 10,000 events per second from sensor and vision sources, with burst capacity to 50,000 events/second.'],
    ['NFR3', 'Availability', 'Production target: 99.9% uptime. Hackathon: single-node recovery within 60 seconds.'],
    ['NFR4', 'Security', 'All API traffic over TLS 1.3. Authentication via JWT with role-based access control (worker, supervisor, admin). Audit log is append-only.'],
    ['NFR5', 'Privacy', 'Camera frames processed ephemerally; no raw video stored without explicit retention policy. Worker profiles respect data minimization.'],
    ['NFR6', 'Extensibility', 'New perception modules (e.g., acoustic anomaly detection) can be added by implementing a single interface and registering with the Event Orchestrator.'],
    ['NFR7', 'Offline Operation', 'Core SHELDRA reasoning functions must operate without internet connectivity. Cloud-dependent features (TTS, embedding) degrade gracefully.'],
    ['NFR8', 'Observability', 'All module decisions logged with trace_id. Prometheus metrics for latency, throughput, error rate per module. Structured logging in JSON format.'],
]
add_table(['ID', 'Requirement', 'Description'], nfr)

# ═══════════════════════════════════════════════════════════════════════════
# 6. OVERALL SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
add_heading('6. Overall System Architecture', level=1)

add_para(
    'The system is organized into five strict layers. Each layer communicates only with its '
    'adjacent layers, enforcing separation of concerns and enabling independent scaling.'
)

add_subsection_heading('6.1 Layer Diagram')
add_code("""
User (Worker / Supervisor / Safety Officer)
     |
     | HTTPS / WSS
     v
+----------------------------------------------------+
|  Presentation Layer                                |
|  Next.js 14 · R3F 3D Avatar · WebRTC Voice       |
|  React Flow (DT Viz) · A-Frame VR                 |
|  Tailwind CSS · shadcn/ui · WebSocket Client      |
+----------------------------------------------------+
     | REST + WebSocket
     v
+----------------------------------------------------+
|  API Layer                                         |
|  FastAPI · Pydantic · WebSocket Manager            |
|  Entry: /api/v1/sheldra/*                          |
|  Auth: JWT + RBAC                                  |
+----------------------------------------------------+
     | Internal IPC
     v
+----------------------------------------------------+
|  SHELDRA Intelligence Engine                       |
|  +------------------+  +------------------+        |
|  | EventOrchestrator|  | LLM Core         |        |
|  | DecisionTree     |  | Personalization  |        |
|  | KG Manager       |  | RAG Engine       |        |
|  | VisionIntel      |  | MonitorIntel     |        |
|  | Memory Manager   |  | Explainability   |        |
|  | Safety Guardrails|  |                  |        |
|  +------------------+  +------------------+        |
+----------------------------------------------------+
     | Kafka / Redis / gRPC
     v
+----------------------------------------------------+
|  Infrastructure Layer                              |
|  Apache Kafka · Redis · MinIO · Service Mesh       |
|  Health Checks · Circuit Breakers · Rate Limiters  |
+----------------------------------------------------+
     | Native drivers
     v
+----------------------------------------------------+
|  Data Layer                                        |
|  Neo4j (Graph) · Qdrant (Vectors)                  |
|  InfluxDB (Time-Series) · PostgreSQL (Relational)  |
+----------------------------------------------------+
     | Managed services
     v
+----------------------------------------------------+
|  Cloud Layer                                       |
|  AWS ECS Fargate · RDS · S3 · CloudFront · Route53 |
+----------------------------------------------------+
""")

add_subsection_heading('6.2 Design Decisions')

decisions = [
    ['Why a single API entry point?', 'All external interactions go through /api/v1/sheldra/*. '
     'This enforces consistent authentication, validation, tracing, and rate limiting. The '
     'engine never exposes internal module endpoints externally.'],
    ['Why strict layering?', 'Each layer can be independently scaled, deployed, tested, and '
     'replaced. The engine has zero knowledge of the frontend. The data layer has zero '
     'knowledge of the business logic.'],
    ['Why Kafka as the nervous system?', 'Kafka provides durable, replayable, partitioned event '
     'storage. Every module that produces events (Vision, Monitoring) writes to Kafka. '
     'Every module that consumes events (Orchestrator, Decision Tree) reads from Kafka. '
     'This decouples producers from consumers in both time and space.'],
]
for title, desc in decisions:
    p = doc.add_paragraph()
    r = p.add_run(title + ' ')
    r.bold = True
    p.add_run(desc)

add_subsection_heading('6.3 Tradeoff Analysis')
tradeoffs = [
    ['Monolithic engine vs. microservices', 'Chosen: Modular engine (single process, module isolation). '
     'Tradeoff: Deploy as one unit = simpler operations, no RPC overhead. Future: Extract high-load '
     'modules (Vision, Monitoring) into separate services with gRPC.'],
    ['Kafka vs. message queue (RabbitMQ)', 'Chosen: Kafka. Tradeoff: Higher operational complexity, '
     'better for replay, audit, and stream processing. RabbitMQ would be simpler for point-to-point '
     'but lacks the durability and partitioning needed for sensor data.'],
    ['Graph DB vs. relational for relationships', 'Chosen: Neo4j (graph). Tradeoff: Graph queries '
     'are 10-100x faster for multi-hop traversals (root cause analysis). Relational would require '
     'costly JOINs. Graph DB adds operational overhead.'],
]
add_table(['Decision', 'Choice', 'Rationale'], tradeoffs)

# ═══════════════════════════════════════════════════════════════════════════
# 7. SHELDRA INTELLIGENCE ENGINE
# ═══════════════════════════════════════════════════════════════════════════
add_heading('7. SHELDRA Intelligence Engine', level=1)

add_para(
    'The SHELDRA Intelligence Engine is a single AI system composed of 11 internal modules. '
    'Every module implements a common interface: `__init__(config)`, `process(context) -> '
    'ModuleResult`, and `health() -> bool`. The EventOrchestrator routes incoming events to '
    'the appropriate modules and assembles their outputs into a unified response.'
)

add_subsection_heading('7.1 Module Directory')
add_code("""
sheldra_engine/
  __init__.py                  # Engine entry point: SHELDRAEngine class
  orchestrator.py              # EventOrchestrator: routing, lifecycle, tracing
  llm_core.py                  # LLM Reasoning Core: prompt assembly, inference, parsing
  decision_tree.py             # Decision Tree Engine: generation, scoring, update cycle
  personalization.py           # Personalization Engine: profiles, adaptation, learning speed
  kg_manager.py                # Knowledge Graph Manager: Cypher builder, CRUD, caching
  rag_engine.py                # RAG Engine: embedding, retrieval, citation tracking
  vision_intel.py              # Vision Intelligence: detection, pose, tracking
  monitoring_intel.py          # Monitoring Intelligence: anomaly, correlation, scoring
  memory_manager.py            # Memory Manager: short-term (Redis), long-term (KG)
  explainability.py            # Explainability Engine: trace, SHAP, confidence
  safety_guardrails.py         # Safety Guardrails: validation, constraints, audit
  models/                      # Pydantic models for inter-module communication
    events.py                  # Event types, payloads, metadata
    context.py                 # Context assembly schemas
    responses.py               # SHELDRA response schema
  tests/                       # Module-level tests
""")

add_subsection_heading('7.2 Module Responsibilities')

modules = [
    ['EventOrchestrator', 'Entry point. Receives all events, validates, routes to modules, '
     'manages lifecycle (validate \u2192 route \u2192 execute \u2192 collect \u2192 respond). '
     'Propagates trace_id. Handles module timeouts and failures.'],
    ['LLM Core', 'Central reasoning. Assembles context from all modules into a prompt. '
     'Executes Llama 3 8B inference. Parses structured output. Supports streaming.'],
    ['DecisionTree', 'Hazard resolution tree management. Generates trees from KG + RAG. '
     'Scores branches on outcomes. Runs 5-min update cycle. Versions and supports rollback.'],
    ['Personalization', 'Worker profile management. Computes adaptation dimensions '
     '(depth, tone, vocabulary, rate, feedback, language). Injects profile into LLM context.'],
    ['KGManager', 'Graph database interface. Parameterized Cypher queries. CRUD for all '
     'node/edge types. Temporal query support. Query caching via Redis.'],
    ['RAGEngine', 'Document retrieval. PDF chunking, E5 embedding, Qdrant hybrid search. '
     'Cross-encoder re-ranking. Citation tracking.'],
    ['VisionIntel', 'Camera perception. YOLOv8n for PPE/person detection. RTMPose for pose. '
     'ByteTrack for ID persistence. Polygon ROI for zone encroachment.'],
    ['MonitorIntel', 'IoT perception. TimesNet anomaly detection. Multi-sensor correlation. '
     'Adaptive thresholding. Alert suppression.'],
    ['MemoryManager', 'State persistence. Short-term: Redis session store with TTL. '
     'Long-term: KG interaction history. Context window assembly for LLM.'],
    ['Explainability', 'Decision tracing. Collects module inputs/outputs. Generates NL '
     'explanations. Computes confidence provenance. SHAP integration.'],
    ['SafetyGuardrails', 'Output validation. Constraint rules. Confidence threshold enforcement. '
     'Procedure contradiction check. Audit log writer.'],
]
add_table(['Module', 'Responsibility'], modules)

add_subsection_heading('7.3 Inter-Module Data Flow')

add_para('Sequence diagram for a SHELDRA chat request:', bold=True)
add_code("""
  Client          API          Orchestrator    LLM Core    KG Mgr   RAG    Personal   Memory
    |              |                |              |          |       |        |         |
    | POST /chat   |                |              |          |       |        |         |
    |------------->|                |              |          |       |        |         |
    |              | validate auth  |              |          |       |        |         |
    |              |----------------|              |          |       |        |         |
    |              | route(event)   |              |          |       |        |         |
    |              |--------------->|              |          |       |        |         |
    |              |                | load_profile |          |       |        |         |
    |              |                |------------------------->|       |        |         |
    |              |                | load_history  |          |       |        |         |
    |              |                |---------------------------------------->|         |
    |              |                | retrieve_know |          |       |        |         |
    |              |                |----------------------->|       |        |         |
    |              |                |---------------------->|       |        |         |
    |              |                | assemble_context      |       |        |         |
    |              |                |--------------->|      |       |        |         |
    |              |                |                | infer |       |        |         |
    |              |                |                |------>|       |        |         |
    |              |                | parse_response |       |       |        |         |
    |              |                |<---------------|      |       |        |         |
    |              |                | validate       |              |        |         |
    |              |                | log_trace      |              |        |         |
    |              |<---------------|                |              |        |         |
    | stream resp  |                |                |              |        |         |
    |<-------------|                |                |              |        |         |
""")

add_subsection_heading('7.4 Event Flow')
add_para(
    'The EventOrchestrator supports multiple event types, each routing to a different set '
    'of modules:'
)
event_flows = [
    ['chat', 'Worker sends message', 'Orchestrator, LLM Core, Personalization, Memory, RAG, KG Manager, Explainability, Guardrails'],
    ['vision.detection', 'Camera detects event', 'Orchestrator, VisionIntel, KG Manager, DecisionTree, LLM Core, Explainability'],
    ['sensor.alert', 'Anomaly detected', 'Orchestrator, MonitorIntel, KG Manager, DecisionTree, LLM Core, Explainability'],
    ['tree.update', 'Decision tree modified', 'Orchestrator, DecisionTree, KG Manager'],
    ['supervisor.action', 'HITL override', 'Orchestrator, KG Manager, Explainability, SafetyGuardrails'],
]
add_table(['Event Type', 'Trigger', 'Modules Invoked'], event_flows)

add_subsection_heading('7.5 Failure Handling')
add_para(
    'Each module has a configurable timeout (default 5s). If a module exceeds the timeout, '
    'the Orchestrator logs the failure (with trace_id), records partial results, and continues '
    'with available data. The LLM Core is informed of which modules failed so it can adjust '
    'its response accordingly (e.g., "I cannot check your current location because the '
    'Knowledge Graph is temporarily unavailable."). The engine degrades gracefully: Vision '
    'failure means no visual feedback, but text chat continues.'
)

add_subsection_heading('7.6 Future Extensibility')
add_para(
    'New modules are added by subclassing the Module base class and registering with the '
    'Orchestrator. Examples: AcousticIntel (microphone array for abnormal sound detection), '
    'AirQualityIntel (particulate matter, VOCs), BiometricIntel (heart rate, skin temperature '
    'from wearables). No existing module needs modification for a new module to be added.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 8. DECISION TREE LEARNING SYSTEM
# ═══════════════════════════════════════════════════════════════════════════
add_heading('8. Decision Tree Learning System', level=1)

add_subsection_heading('8.1 Why')
add_para(
    'Static safety checklists fail because they cannot adapt to new information. If a checklist '
    'step proves ineffective, it remains in the procedure until a human updates it. The Decision '
    'Tree Learning System replaces static checklists with self-improving graphs that learn from '
    'every incident outcome.'
)

add_subsection_heading('8.2 Tree Data Model')
add_code("""
{
  "tree_id": "gas-leak-response",
  "hazard_type": "gas_leak",
  "version": 7,
  "created_at": "2026-07-07T08:00:00Z",
  "last_updated": "2026-07-07T13:30:00Z",
  "overall_accuracy": 0.89,
  "total_evaluations": 152,
  "nodes": [
    {
      "node_id": "n001",
      "type": "question",
      "content": "Gas level > 50ppm?",
      "conditions": { "sensor.gas_ppm": { "gt": 50 } },
      "branches": [
        { "to": "n002", "label": "Yes", "action": "evacuate_zone" },
        { "to": "n003", "label": "No", "action": "increase_monitoring" }
      ],
      "accuracy": 0.94,
      "times_taken": 47
    },
    {
      "node_id": "n002",
      "type": "action",
      "content": "Initiate Zone B evacuation",
      "branches": [
        { "to": "n004", "label": "Evacuated", "action": "seal_zone" },
        { "to": "n005", "label": "Remaining", "action": "alert_rescue" }
      ],
      "accuracy": 0.91,
      "times_taken": 23
    }
  ],
  "root_id": "n001"
}
""")

add_subsection_heading('8.3 Generation Pipeline')
add_para(
    'When a new hazard type is encountered, the system generates an initial tree from three sources:'
)
add_bullet('Knowledge Graph: Hazard taxonomy nodes link to applicable SafetyRule nodes. Each rule contributes a decision node with conditions and branches.')
add_bullet('RAG Engine: Procedure documents for the hazard type are parsed into action nodes. Steps become sequential action nodes with confirmation branches.')
add_bullet('Historical Incidents: Past incidents of the same hazard type are analyzed for common resolution paths. Frequent paths become preferred branches with higher initial confidence.')
add_para('')
add_para('The generator produces a directed acyclic graph validated for completeness: all branches must terminate in a resolution or escalation node.', bold=False)

add_subsection_heading('8.4 Update Cycle (5 Minutes)')
add_para(
    'A background daemon executes every 5 minutes. Each cycle:'
)
add_bullet('Collect Outcomes: Query the Knowledge Graph for all completed SHELDRASession nodes where the hazard matches this tree and the session completed in the last 5 minutes.')
add_bullet('Score Each Branch: For each session, compare the predicted resolution path against the actual outcome. Success=1.0, Partial=0.5, Failure=0.0. Update the rolling 10-trial accuracy for each traversed node.')
add_bullet('Prune: Any branch with accuracy < 0.5 over the last 10 evaluations is flagged. If it remains below 0.5 for 3 consecutive cycles, the branch is removed.')
add_bullet('Generate: Analyze failed resolutions for common patterns. If a pattern appears 3+ times, propose a new branch with the alternative action.')
add_bullet('Version: Increment tree version. Store the previous version in the Knowledge Graph with its accuracy metrics for rollback.')
add_bullet('Notify: Write a tree.update event to Kafka. SHELDRA can inform supervisors of the update.')

add_subsection_heading('8.5 Accuracy Measurement')
acc_table = [
    ['Success', '1.0', 'Hazard resolved correctly. No injuries. No escalation needed. Worker followed recommended path.'],
    ['Partial', '0.5', 'Hazard resolved but with delay, confusion, or deviation. Escalation not required but supervisor noted issues.'],
    ['Failure', '0.0', 'Escalation required. Incorrect procedure followed. Near-miss or injury occurred.'],
]
add_table(['Outcome', 'Score', 'Definition'], acc_table)
add_para('')
add_para(
    'Node accuracy = sum(scores) / count over rolling 10-trial window. Tree accuracy = '
    'mean of all leaf node accuracies. Displayed on the dashboard with trend line.'
)

add_subsection_heading('8.6 Technology Choice')
add_para(
    'The Decision Tree Engine is implemented in pure Python without a workflow engine '
    '(e.g., Camunda, Temporal) because: (a) trees are small data structures (10-50 nodes) '
    'that fit in memory, (b) the update cycle is computation-light (score aggregation, '
    'not heavy processing), and (c) eliminating the workflow engine reduces deployment '
    'complexity for the hackathon. Future enterprise deployments may adopt Temporal for '
    'long-running workflow orchestration.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 9. COMPUTER VISION PIPELINE
# ═══════════════════════════════════════════════════════════════════════════
add_heading('9. Computer Vision Pipeline', level=1)

add_subsection_heading('9.1 Why')
add_para(
    'Camera feeds are the richest asynchronous data source in industrial environments. The '
    'Vision Intelligence module transforms raw frames into structured observations that the '
    'Event Orchestrator can route to the LLM Core for SHELDRA feedback.'
)

add_subsection_heading('9.2 Pipeline Architecture')
add_code("""
RTSP Stream
    |
    v
FFmpeg Frame Grabber (configurable FPS, default 5)
    |
    v
Frame Buffer (bounded queue, max 100 frames)
    |
    v
+------------------------------------------------+
| Parallel Inference                              |
| YOLOv8n (PPE + Person)    RTMPose (Pose)        |
| 6 classes, nms=0.5        17 keypoints          |
|    |                           |                |
|    v                           v                |
| ByteTrack (person ID)    Behavior Analyzer      |
| IoU matching across      Reaching, climbing,    |
| frames for persistence   fall detection         |
|    |                           |                |
+----+---------------------------+----------------+
    |
    v
Structured Event:
{
  "event_type": "vision.detection",
  "detections": [
    { "class": "missing_helmet", "worker_id": "w3124",
      "zone": "A3", "confidence": 0.94, "bbox": [...] }
  ],
  "timestamp": "...",
  "camera_id": "cam-07"
}
    |
    v
Kafka Topic: vision.detections
    |
    v
EventOrchestrator -> modules...
""")

add_subsection_heading('9.3 Models and Performance')
cv_perf = [
    ['PPE Detection', 'YOLOv8n (nano)', 'CPU-only', '~15ms', '30+ FPS on CPU via ONNX. Detects all 6 PPE classes in single forward pass.'],
    ['Person Detection', 'YOLOv8s (small)', 'GPU (optional)', '2-3ms on T4', 'Larger model for accurate person detection. Runs on GPU if available, ONNX CPU otherwise.'],
    ['Person Tracking', 'ByteTrack', 'CPU', '<1ms', 'IoU-based association across frames. Maintains worker IDs for SHELDRA to address specific workers.'],
    ['Pose Estimation', 'RTMPose-m', 'GPU', '~5ms/person', '17-keypoint pose. Enables ergonomic feedback and fall detection.'],
    ['Zone Encroachment', 'Polygon ROI + IoU', 'CPU', '<1ms', 'No ML needed. Pre-defined zone polygons. Intersection of person bbox with restricted zone.'],
]
add_table(['Task', 'Model', 'Hardware', 'Latency', 'Details'], cv_perf)

add_subsection_heading('9.4 Data Flow')
add_para(
    'Each camera runs as an independent processing thread. Frame grabbing and inference are '
    'pipelined: while frame N is being inferred, frame N+1 is being grabbed. The frame buffer '
    'provides backpressure: if inference is slower than the grab rate, old frames are dropped '
    '(newest-first eviction). This ensures the system always processes the most recent frame, '
    'not a backlog.'
)

add_subsection_heading('9.5 Technology Choice')
add_para(
    'ONNX Runtime was chosen over TensorRT and PyTorch JIT because: (a) ONNX provides a single '
    'runtime for all models (vision + time-series), (b) ONNX supports both CPU and GPU with '
    'the same API, and (c) ONNX execution providers (TensorRT, OpenVINO, CoreML) allow '
    'hardware-specific optimization without code changes. The tradeoff is that ONNX may be '
    '5-10% slower than a native TensorRT deployment for GPU inference.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 10. TIME SERIES ANALYTICS
# ═══════════════════════════════════════════════════════════════════════════
add_heading('10. Time Series Analytics', level=1)

add_subsection_heading('10.1 Why')
add_para(
    'IoT sensors produce continuous numerical data that must be analyzed for anomalies. The '
    'Monitoring Intelligence module detects deviations from expected patterns and produces '
    'hazard triggers that feed into the Decision Tree system.'
)

add_subsection_heading('10.2 Anomaly Detection Pipeline')
add_code("""
Sensor Network (simulated)
    |
    | MQTT / Kafka
    v
Kafka Topic: sensor.readings
    |
    +---- Partition by sensor_id (temp_01, vib_02, gas_03)
    |
    v
+----------------------------------------------------+
| Sliding Window Buffer (60 readings per sensor)     |
| Z-score normalization per sensor type              |
| Linear interpolation for < 5s gaps                 |
+----------------------------------------------------+
    |
    v
+----------------------------------------------------+
| TimesNet Inference                                  |
| Input: (60, num_sensors) tensor                    |
| Output: reconstruction error per sensor            |
| Anomaly = reconstruction_error > adaptive_thresh   |
+----------------------------------------------------+
    |
    v
+----------------------------------------------------+
| Multi-Sensor Correlation                            |
| Weighted compound score:                            |
|   Vib(0.4) + Temp(0.3) + Gas(0.3)                 |
| Compound > 0.7 -> compound_alert                   |
+----------------------------------------------------+
    |
    v
Kafka Topic: system.alerts
    |
    v
EventOrchestrator -> DecisionTree -> SHELDRA
""")

add_subsection_heading('10.3 Model Details')
add_para(
    'TimesNet (Microsoft Research, ICLR 2023 Best Paper) transforms 1D time series into 2D '
    'tensors using Fast Fourier Transform for period detection. This captures multiple '
    'temporal patterns simultaneously: shift patterns (8-hour cycles), machine cycles '
    '(varying by equipment type), and seasonal variations (day/night, weekday/weekend). '
    'The reconstruction error is the anomaly score: high error means the current pattern '
    'does not match learned normal patterns.'
)

add_subsection_heading('10.4 Adaptive Thresholding')
add_para(
    'Anomaly thresholds are not static. The system maintains a rolling baseline per '
    'sensor per time-of-day bucket (1-hour windows). A "spike" at 3 AM that is normal '
    'for shift change at 7 AM would not trigger an alert. This reduces false positives '
    'by approximately 60% compared to fixed thresholds.'
)

add_subsection_heading('10.5 Alert Suppression')
add_para(
    'A 30-second debounce window prevents alert storms. If the same sensor triggers an '
    'anomaly every second for 10 seconds, only one alert is emitted. The alert includes '
    'the peak anomaly score and duration. This is critical for industrial environments '
    'where sensors can oscillate near a threshold.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 11. KNOWLEDGE GRAPH
# ═══════════════════════════════════════════════════════════════════════════
add_heading('11. Knowledge Graph', level=1)

add_subsection_heading('11.1 Why')
add_para(
    'Relational databases become inefficient when querying multi-hop relationships: "Find '
    'all workers who are in Zone A, who have been involved in incidents similar to the '
    'current one, who speak Spanish, and who have a supervisor currently on shift." '
    'A graph database traverses these relationships in constant time per hop, whereas '
    'a relational database requires multiple JOIN operations that scale with table size.'
)

add_subsection_heading('11.2 Schema (Nodes)')

nodes_schema = [
    ['Worker', 'worker_id, name, role, experience_level, learning_speed, language, certifications[], ppe_compliance_rate, emotional_state'],
    ['Hazard', 'hazard_id, type, severity, frequency, first_observed, last_observed, tree_ids[]'],
    ['DecisionTree', 'tree_id, hazard_type, version, current, accuracy, last_updated, root_node_id'],
    ['Incident', 'incident_id, type, severity, status, timestamp, resolution_path[], accuracy_score'],
    ['SafetyRule', 'rule_id, description, regulation_ref, violation_conditions, severity, action_required'],
    ['Procedure', 'procedure_id, title, hazard_types[], steps[], required_ppe[], version'],
    ['SHELDRASession', 'session_id, worker_id, hazard_id, tree_id, messages[], outcome, accuracy_score, start_time, end_time'],
    ['Zone', 'zone_id, name, type, hazard_level, required_ppe[], active_hazards[]'],
    ['Equipment', 'equipment_id, name, type, status, last_inspection, maintenance_due, associated_hazards[]'],
    ['Supervisor', 'supervisor_id, name, role, active, escalation_level'],
]
add_table(['Node Label', 'Key Properties'], nodes_schema)

add_subsection_heading('11.3 Schema (Edges)')

edges_schema = [
    ['HAS_PROFILE', 'Worker \u2192 Worker', 'Links worker to their profile properties'],
    ['GUIDED_BY', 'Worker \u2192 SHELDRASession', 'Records every SHELDRA coaching interaction with a worker'],
    ['ABOUT', 'SHELDRASession \u2192 Hazard', 'Associates a coaching session with a specific hazard'],
    ['RESOLVED_VIA', 'Incident \u2192 DecisionTree', 'Records which decision tree version was used for an incident'],
    ['REQUIRES', 'Zone \u2192 SafetyRule', 'Safety rules that apply to a specific zone'],
    ['VIOLATES', 'Observation \u2192 SafetyRule', 'Links a detected violation to the rule that was broken'],
    ['TRIGGERED_BY', 'Incident \u2192 Observation', 'Root cause: which sensor reading or vision detection triggered this'],
    ['LOCATED_IN', 'Worker/Equipment \u2192 Zone', 'Current location of an entity'],
    ['MONITORED_BY', 'Zone \u2192 Camera/Sensor', 'Which monitoring devices cover this zone'],
    ['MANAGES', 'Supervisor \u2192 Worker', 'Supervisor-worker reporting relationship'],
    ['SIMILAR_TO', 'Incident \u2192 Incident', 'Similar incident based on graph edit distance or feature similarity'],
]
add_table(['Edge Type', 'Source \u2192 Target', 'Semantics'], edges_schema)

add_subsection_heading('11.4 Key Queries')
add_code("""
// Find all workers in Zone A with their PPE compliance
MATCH (w:Worker)-[:LOCATED_IN]->(z:Zone {zone_id: 'A'})
RETURN w.worker_id, w.name, w.ppe_compliance_rate
ORDER BY w.ppe_compliance_rate ASC

// Root cause analysis: trace incident to its trigger
MATCH path = (i:Incident {incident_id: $id})-[:TRIGGERED_BY*]->(obs)
RETURN nodes(path)

// Worker interaction history for personalization
MATCH (w:Worker {worker_id: $id})-[:GUIDED_BY]->(s:SHELDRASession)
RETURN s.hazard_id, s.outcome, s.accuracy_score, s.start_time
ORDER BY s.start_time DESC
LIMIT 20

// Most effective decision tree for a hazard
MATCH (i:Incident {hazard_type: $type})-[:RESOLVED_VIA]->(dt:DecisionTree)
RETURN dt.tree_id, dt.version, avg(i.accuracy_score) as avg_accuracy
ORDER BY avg_accuracy DESC
LIMIT 1
""")

add_subsection_heading('11.5 Technology Choice')
add_para(
    'Neo4j was chosen over Amazon Neptune and ArangoDB because: (a) Cypher is the most '
    'intuitive graph query language for developers, (b) Neo4j\'s property graph model maps '
    'directly to industrial entities (nodes have rich properties, edges carry weight and '
    'temporal metadata), (c) the Python driver (neo4j) is mature and well-documented, and '
    '(d) Neo4j AuraDB provides managed cloud deployment. The tradeoff is that Neo4j is '
    'not horizontally sharded by default; enterprise deployments require Neo4j Fabric.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 12. RETRIEVAL AUGMENTED GENERATION
# ═══════════════════════════════════════════════════════════════════════════
add_heading('12. Retrieval Augmented Generation', level=1)

add_subsection_heading('12.1 Why')
add_para(
    'Large Language Models generate plausible but factually incorrect text (hallucination). '
    'In industrial safety, a hallucinated procedure can cause injury. RAG grounds every '
    'SHELDRA response in retrievable safety documentation, providing citations for every claim.'
)

add_subsection_heading('12.2 Architecture')
add_code("""
Document Source (PDF)
    |
    v
+------------------------------------+
| Ingestion Pipeline                 |
| pdfplumber -> text extraction      |
| RecursiveCharacterTextSplitter     |
| chunk_size=512, overlap=50         |
| Preserves section headers, pages   |
+------------------------------------+
    |
    v
+------------------------------------+
| Embedding                          |
| E5-mistral-7b -> 4096-dim vector   |
| ONNX Runtime (FP16 quantization)   |
| Per chunk, stored in Qdrant        |
+------------------------------------+
    |
    v
+------------------------------------+
| Qdrant Collection                  |
| Hybrid: dense (cosine) + sparse    |
| (BM25). Payload: doc_id, section,  |
| page, chunk_index, source_file     |
+------------------------------------+
    |
    v
+------------------------------------+
| Retrieval (per incoming query)     |
| Query -> Qdrant hybrid search      |
| Top-20 results -> Cross-encoder    |
| Re-rank -> Top-3                   |
+------------------------------------+
    |
    v
+------------------------------------+
| Context Assembly                   |
| Retrieved chunks + source citations|
| Injected into LLM prompt           |
+------------------------------------+
    |
    v
LLM Core -> Generated Response with "[Source: OSHA 1910.134 Section 4.2]"
""")

add_subsection_heading('12.3 Retrieval Strategy')
add_para(
    'Hybrid search combines semantic similarity (dense vector search) with keyword matching '
    '(sparse BM25). This is essential for industrial safety queries where exact terminology '
    'matters: "OSHA 1910.134" must match the specific regulation number, not a semantically '
    'similar phrase. Dense search captures conceptual queries like "what to do during a gas '
    'leak." The cross-encoder re-ranker (a smaller BERT model) scores the top-20 results '
    'and returns the top-3 most relevant chunks.'
)

add_subsection_heading('12.4 Citation Format')
add_para(
    'Every SHELDRA response that references procedure includes a citation block: "Per OSHA '
    '1910.134 Section 4.2: \'Respirators must be fitted before entering confined spaces.\'" '
    'If the retrieval confidence score is below 0.5, SHELDRA responds with: "I cannot find '
    'a specific procedure for this situation. I recommend consulting your supervisor."'
)

# ═══════════════════════════════════════════════════════════════════════════
# 13. DATA FLOW ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
add_heading('13. Data Flow Architecture', level=1)

add_subsection_heading('13.1 Event Taxonomy')

event_tax = [
    ['sheldra.interactions', 'Chat messages, SHELDRA responses', 'API \u2192 Engine \u2192 API', 'Audit, training, tree improvement'],
    ['vision.detections', 'PPE violations, zone encroachments, behavior alerts', 'VisionIntel \u2192 Kafka \u2192 Engine', 'SHELDRA feedback triggers'],
    ['sensor.readings', 'Raw IoT telemetry (temp, vib, gas, noise)', 'Simulator/Sensors \u2192 Kafka \u2192 MonitorIntel', 'Anomaly detection input'],
    ['system.alerts', 'Anomaly alerts, compound risk alerts, system health', 'MonitorIntel/Engine \u2192 Kafka \u2192 API', 'Dashboard display, SHELDRA awareness'],
    ['decision.trees', 'Tree updates, version changes, accuracy scores', 'DecisionTree \u2192 Kafka \u2192 KG', 'Tree persistence, frontend updates'],
    ['worker.profiles', 'Profile updates, learning speed changes, training scores', 'Engine \u2192 Kafka \u2192 KG', 'Personalization data durability'],
    ['supervisor.actions', 'Overrides, acknowledgments, dismissals, escalations', 'API \u2192 Kafka \u2192 Engine', 'HITL audit trail'],
]
add_table(['Topic', 'Content', 'Flow', 'Purpose'], event_tax)

add_subsection_heading('13.2 Message Schema (Avro)')
add_code("""
{
  "namespace": "com.sheldra.event",
  "type": "record",
  "name": "SHELDRAEvent",
  "fields": [
    { "name": "event_id",         "type": "string" },
    { "name": "event_type",       "type": "string" },
    { "name": "trace_id",         "type": "string" },
    { "name": "source_module",    "type": "string" },
    { "name": "timestamp",        "type": "long"   },
    { "name": "worker_id",        "type": ["null", "string"], "default": null },
    { "name": "payload",          "type": "bytes"  },
    { "name": "metadata",         "type": { "type": "map", "values": "string" } }
  ]
}
""")

add_subsection_heading('13.3 Event Flow for a Typical SHELDRA Interaction')
add_code("""
Worker: "What should I do about the vibration on machine #7?"
    |
    | 1. POST /api/v1/sheldra/chat
    v
API Layer -> validate auth -> get worker profile -> create trace_id
    |
    | 2. Route to EventOrchestrator
    v
Orchestrator:
    |-- 3. MemoryManager.get_short_term(trace_id)  [Redis]
    |-- 4. MemoryManager.get_long_term(worker_id)   [KG query]
    |-- 5. PersonalizationEngine.load_profile(worker_id)
    |-- 6. RAGEngine.retrieve("vibration machine #7 procedure")
    |-- 7. KGManager.query("equipment status machine #7")
    |-- 8. KGManager.query("recent incidents machine #7")
    |
    | 9. Assemble context -> LLM Core
    v
LLM Core:
    |-- 10. Build prompt (system + profile + context + history)
    |-- 11. Llama 3 inference -> structured response
    |-- 12. Parse {message, tone, gesture, confidence, citations}
    |
    | 13. SafetyGuardrails.validate(response)
    | 14. ExplainabilityEngine.record_trace(module_inputs, outputs)
    | 15. MemoryManager.update_short_term(trace_id, conversation)
    | 16. KGManager.create_session_node(worker_id, hazard, response)
    |
    | 17. Return response to API
    v
API Layer -> stream to client
    |
    | 18. SHELDRA avatar speaks, gestures, displays message
    v
Worker hears SHELDRA's response
""")

# ═══════════════════════════════════════════════════════════════════════════
# 14. DATABASE DESIGN
# ═══════════════════════════════════════════════════════════════════════════
add_heading('14. Database Design', level=1)

add_subsection_heading('14.1 Polyglot Persistence Model')
db_design = [
    ['Neo4j', 'Graph', 'Entity relationships, worker profiles, incident graphs, decision trees, session logs', 'Multi-hop queries, root cause analysis, pattern matching', 'Temporal edges for point-in-time queries; query caching via Redis'],
    ['Qdrant', 'Vector', 'Document embeddings for RAG, feature vectors for similarity search', 'Semantic search, hybrid (dense+sparse) retrieval', 'Payload filtering for metadata; quantization for memory efficiency'],
    ['InfluxDB', 'Time-Series', 'IoT sensor readings, model inference metrics', 'Windowed aggregates, continuous queries, downsampling', 'Retention policies; 7-day raw, 1-year downsampled'],
    ['PostgreSQL', 'Relational', 'Users, teams, auth tokens, configuration, audit log', 'Transactional CRUD, reporting, auth', 'Row-level security for multi-tenant isolation'],
    ['Redis', 'Key-Value', 'Session state, conversation buffers, KG query cache, rate limiter counters', 'Sub-millisecond reads, Pub/Sub, TTL-based expiry', 'LRU eviction; persistence disabled (cache only)'],
    ['MinIO', 'Object', 'Camera frames (transient), model artifacts, report PDFs, training data', 'Per-frame storage for incident evidence, batch processing', 'S3-compatible; lifecycle policies for automatic tiering'],
]
add_table(['Store', 'Type', 'Data', 'Access Pattern', 'Configuration'], db_design)

add_subsection_heading('14.2 PostgreSQL Schema (Relational Data)')
add_code("""
-- Users and Authentication
CREATE TABLE users (
    user_id     UUID PRIMARY KEY DEFAULT gen_random_uuid(),
    email       VARCHAR(255) UNIQUE NOT NULL,
    role        VARCHAR(50) NOT NULL CHECK (role IN ('worker', 'supervisor', 'admin')),
    password_hash VARCHAR(255) NOT NULL,
    worker_id   VARCHAR(50) REFERENCES workers(worker_id),
    created_at  TIMESTAMPTZ DEFAULT NOW(),
    last_login  TIMESTAMPTZ
);

-- Audit Log (append-only)
CREATE TABLE audit_log (
    log_id      BIGSERIAL PRIMARY KEY,
    trace_id    VARCHAR(64) NOT NULL,
    event_type  VARCHAR(50) NOT NULL,
    actor_id    VARCHAR(50),
    action      VARCHAR(100) NOT NULL,
    details     JSONB,
    created_at  TIMESTAMPTZ DEFAULT NOW()
);
CREATE INDEX idx_audit_trace ON audit_log(trace_id);
CREATE INDEX idx_audit_time ON audit_log(created_at DESC);

-- Configuration
CREATE TABLE config (
    config_key   VARCHAR(100) PRIMARY KEY,
    config_value JSONB NOT NULL,
    updated_at   TIMESTAMPTZ DEFAULT NOW()
);
""")

# ═══════════════════════════════════════════════════════════════════════════
# 15. API ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
add_heading('15. API Architecture', level=1)

add_subsection_heading('15.1 Design Principles')
add_bullet('Single entry point: All external interactions through /api/v1/sheldra/*')
add_bullet('REST for synchronous operations (profile fetch, tree history)')
add_bullet('WebSocket for streaming (SHELDRA responses, real-time updates)')
add_bullet('Stateless API servers (session state in Redis, not in memory)')
add_bullet('Versioned from day one (v1 prefix)')

add_subsection_heading('15.2 Core Endpoints')
api_eps = [
    ['POST', '/api/v1/sheldra/chat', 'JSON: {worker_id, message, context, session_id}', 'Streaming response (SSE/WebSocket)', 'Primary SHELDRA interaction'],
    ['WS', '/ws/sheldra/stream', 'Auth token', 'JSON messages', 'Real-time bidirectional streaming'],
    ['GET', '/api/v1/sheldra/trace/{trace_id}', 'Path: trace_id', 'JSON: trace object', 'Decision trace retrieval'],
    ['GET', '/api/v1/decision-trees/{hazard_id}', 'Path: hazard_id', 'JSON: current tree', 'Flow chart rendering'],
    ['GET', '/api/v1/decision-trees/{id}/history', 'Path: tree_id', 'JSON: version array', 'Version comparison'],
    ['GET', '/api/v1/workers/{id}/profile', 'Path: worker_id', 'JSON: profile', 'Personalization display'],
    ['PUT', '/api/v1/workers/{id}/profile', 'JSON: profile fields', 'JSON: updated profile', 'Profile management'],
    ['POST', '/api/v1/alerts/{id}/acknowledge', 'JSON: {supervisor_id}', 'JSON: {status}', 'HITL acknowledgment'],
    ['POST', '/api/v1/alerts/{id}/override', 'JSON: {reason, action}', 'JSON: {status}', 'HITL override'],
    ['GET', '/api/v1/analytics/accuracy', 'Query: ?hazard_type', 'JSON: accuracy metrics', 'Accuracy dashboard'],
]
add_table(['Method', 'Path', 'Request', 'Response', 'Purpose'], api_eps)

add_subsection_heading('15.3 Example Request/Response')
add_code("""
POST /api/v1/sheldra/chat
Authorization: Bearer <jwt>
Content-Type: application/json

{
  "worker_id": "w3124",
  "message": "What should I do about the vibration on machine #7?",
  "context": {
    "location": "Zone B",
    "shift": "day",
    "language": "en"
  },
  "session_id": "sess-abc-123"
}

Response (SSE stream):
data: {"token": "I", "done": false}
data: {"token": " see", "done": false}
data: {"token": " that", "done": false}
data: {"token": " machine", "done": false}
data: {"token": " #7", "done": false}
data: {"token": " has", "done": false}
data: {"token": " elevated", "done": false}
data: {"token": " vibration", "done": false}
data: {"token": ". Let me check...", "done": false}
...
data: {
  "done": true,
  "full_response": "I see that machine #7 has elevated vibration readings. "
                   "Per the maintenance procedure (Section 4.2): "
                   "1. Power down the machine using the emergency stop. "
                   "2. Contact maintenance team on channel 3. "
                   "3. Do not restart until maintenance confirms safety.",
  "tone": "calm",
  "gesture": "pointing_left",
  "confidence": 0.89,
  "citations": [
    {"doc": "Maintenance_SOP_v3.pdf", "section": "4.2", "page": 12}
  ],
  "trace_id": "trace-xyz-789"
}
""")

# ═══════════════════════════════════════════════════════════════════════════
# 16. FRONTEND ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
add_heading('16. Frontend Architecture', level=1)

add_subsection_heading('16.1 Directory Structure')
add_code("""
frontend/
  app/                          # Next.js 14 App Router
    layout.tsx                  # Root layout (theme provider, font)
    page.tsx                    # SHELDRA Command Center (default view)
    supervisor/                 # Supervisor console route
    analytics/                  # Accuracy dashboard route
    training/                   # VR training route
    workers/[id]/               # Worker profile route
  components/
    sheldra/
      Avatar.tsx                # R3F 3D avatar component
      ConversationPanel.tsx     # Streaming message display
      SpeechBubble.tsx          # Avatar speech bubble overlay
      ContextPanel.tsx          # Current context display
      ThinkingIndicator.tsx     # "SHELDRA is thinking..." animation
    decision-tree/
      FlowChart.tsx             # React Flow DAG renderer
      NodeDetail.tsx            # Click-to-expand node detail
      StepWalkthrough.tsx       # Step-by-step guided mode
      VersionBadge.tsx          # Version indicator
    dashboard/
      Layout.tsx                # Three-column responsive layout
      AlertQueue.tsx            # Supervisor alert list
      InteractionFeed.tsx       # Live SHELDRA activity log
      AccuracyChart.tsx         # Tree accuracy trend chart
      WorkerCard.tsx            # Worker profile summary card
    voice/
      MicrophoneButton.tsx      # Push-to-talk button
      VoiceWaveform.tsx         # Audio visualization
      CaptionOverlay.tsx        # Text captions for audio
    vr/
      VRScene.tsx               # A-Frame / Three.js XR scene
      VRHazardSimulation.tsx    # Hazard trigger in VR
      VRScoreDisplay.tsx        # Training score overlay
    shared/
      WebSocketProvider.tsx     # WebSocket connection context
      AuthProvider.tsx          # JWT auth context
      ThemeProvider.tsx         # Dark/light mode
  hooks/
    useWebSocket.ts             # WebSocket connection hook
    useSHELDRA.ts               # SHELDRA chat API hook
    useWorkerProfile.ts         # Worker profile query hook
    useDecisionTree.ts          # Decision tree query hook
  lib/
    api.ts                      # API client (fetch + WebSocket)
    types.ts                    # TypeScript type definitions
    constants.ts                # API URLs, config constants
""")

add_subsection_heading('16.2 Component Architecture')
add_para(
    'The frontend follows a presenter pattern: components render data but never compute or '
    'transform it. All business logic resides in the SHELDRA Intelligence Engine. The '
    'WebSocketProvider maintains a persistent connection to /ws/sheldra/stream and dispatches '
    'events to registered components. The useSHELDRA hook encapsulates the chat API, '
    'managing request state, streaming response parsing, and error recovery.'
)

add_subsection_heading('16.3 State Management')
add_para(
    'State is managed through React Context + SWR for cache synchronization. The WebSocket '
    'Provider is the single source of truth for real-time state. Components subscribe to '
    'event types (new_alert, tree_update, worker_status_change) and update local state '
    'accordingly. There is no global state store (Redux, Zustand) because the component '
    'tree is shallow (3 levels max) and data flows are unidirectional: WebSocket \u2192 '
    'Provider \u2192 Hook \u2192 Component.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 17. SECURITY
# ═══════════════════════════════════════════════════════════════════════════
add_heading('17. Security', level=1)

sec_items = [
    ['Authentication', 'JWT-based. Tokens issued on login, valid for 8 hours (shift duration). Refresh tokens with 30-day validity for mobile apps.'],
    ['Authorization', 'Role-based (worker, supervisor, admin). Workers can only read their own profile and interactions. Supervisors can read all profiles in their zone. Admins have full access.'],
    ['Transport', 'TLS 1.3 for all HTTP and WebSocket connections. HSTS headers. Certificate auto-renewal via Let\'s Encrypt.'],
    ['API Security', 'Rate limiting: 30 requests/second per user for chat endpoint. Input validation via Pydantic (type checking, length limits, injection prevention).'],
    ['Audit', 'All HITL actions, profile changes, and tree updates are written to the append-only audit_log table. Audit log has its own database user with INSERT-only grants.'],
    ['Data Privacy', 'Camera frames processed in memory; raw frames written to MinIO only for incident evidence (configurable, default off). Worker profiles contain only safety-relevant data (not HR data).'],
    ['Dependency Scanning', 'GitHub Dependabot for vulnerability alerts. Weekly `npm audit` and `pip audit` in CI pipeline.'],
]
for name, desc in sec_items:
    p = doc.add_paragraph()
    r = p.add_run(name + ': ')
    r.bold = True
    p.add_run(desc)

# ═══════════════════════════════════════════════════════════════════════════
# 18. DEPLOYMENT ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
add_heading('18. Deployment Architecture', level=1)

add_subsection_heading('18.1 Hackathon Deployment (Single Node)')
add_code("""
+--------------------------------------------------------------+
| Single Machine (16GB RAM, 4+ cores, optional NVIDIA GPU)     |
|                                                              |
| docker-compose.yml:                                          |
|   - fastapi (sheldra api)                                    |
|   - sheldra-engine (11 modules)                              |
|   - nextjs (frontend + avatar)                               |
|   - ollama (llama 3 8b)                                      |
|   - zookeeper + kafka                                        |
|   - neo4j                                                     |
|   - qdrant                                                    |
|   - influxdb                                                  |
|   - postgres                                                  |
|   - redis                                                     |
|   - minio                                                     |
|   - nginx (reverse proxy + ssl)                              |
+--------------------------------------------------------------+
""")

add_subsection_heading('18.2 Production Deployment')
add_code("""
+--------------------------------------------------+
| AWS Cloud                                         |
|                                                   |
| VPC                                               |
| +--------------------------------------------+   |
| | Public Subnet                               |   |
| | CloudFront (CDN) -> ALB -> ECS Fargate      |   |
| |   frontend (Next.js, 2x tasks, auto-scale)  |   |
| +--------------------------------------------+   |
|                                                   |
| +--------------------------------------------+   |
| | Private Subnet (App)                       |   |
| | ECS Fargate:                                |   |
| |   sheldra-api (FastAPI, 2x tasks)          |   |
| |   sheldra-engine (4x tasks, CPU)           |   |
| |   vision-inference (2x tasks, GPU if avail)|   |
| |   monitor-inference (2x tasks, CPU)        |   |
| |   decision-tree-worker (1x task, cron)     |   |
| |   ollama (1x task, GPU if avail)           |   |
| +--------------------------------------------+   |
|                                                   |
| +--------------------------------------------+   |
| | Private Subnet (Data)                      |   |
| | Amazon MSK (Kafka)                          |   |
| | RDS PostgreSQL (Multi-AZ)                  |   |
| | ElastiCache Redis (Cluster mode)           |   |
| | Neo4j AuraDB (Managed)                     |   |
| | Qdrant Cloud (Managed)                     |   |
| | InfluxDB Cloud (Managed)                   |   |
| | S3 (object storage)                        |   |
| +--------------------------------------------+   |
+--------------------------------------------------+
""")

add_subsection_heading('18.3 CI/CD Pipeline')
add_para(
    'GitHub Actions workflow: On PR to main, run lint + typecheck + unit tests. On merge to '
    'main, build Docker images, push to ECR, deploy to ECS Fargate (blue/green deployment). '
    'Database migrations run as a separate step before new task set activates.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 19. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════════════════
add_heading('19. Technology Stack', level=1)

stack = [
    ['LLM Reasoning', 'Llama 3 8B (Meta)', 'Open-weight, offline-capable, 8B fits single GPU', 'Ollama (local) -> Together AI (prod)'],
    ['Text Embeddings', 'E5-mistral-7b (Microsoft)', 'Best-in-class quality, 4096-dim, multilingual', 'ONNX Runtime (FP16 quantized)'],
    ['Object Detection', 'YOLOv8n/s (Ultralytics)', 'Real-time, 6-class PPE, fine-tunable', 'ONNX Runtime (CPU/GPU)'],
    ['Pose Estimation', 'RTMPose-m (OpenMMLab)', '17-keypoint, 5ms/person on GPU', 'ONNX Runtime'],
    ['Time-Series', 'TimesNet (Microsoft Research)', 'ICLR 2023 Best Paper, multi-period capture', 'ONNX Runtime'],
    ['API Framework', 'FastAPI (Python)', 'Async-native, auto-docs, Pydantic validation', 'Uvicorn + Gunicorn'],
    ['Frontend', 'Next.js 14 + TypeScript', 'SSR/SSG, React Server Components, App Router', 'Vercel / ECS'],
    ['3D Rendering', 'React Three Fiber + Three.js', 'Declarative 3D in React, WebGL', 'Browser native'],
    ['VR', 'A-Frame (Mozilla)', 'WebXR, no app store, desktop fallback', 'Browser native'],
    ['Message Broker', 'Apache Kafka + Avro', 'Durable, partitioned, replayable', 'Confluent Cloud / MSK'],
    ['Stream Processing', 'KSQL / Flink', 'SQL-level CEP; Flink for complex patterns', 'Confluent Cloud'],
    ['Graph Database', 'Neo4j + Cypher', 'Mature, ACID, Python driver', 'Neo4j AuraDB'],
    ['Vector Database', 'Qdrant', 'Fastest (Rust), hybrid search, payload filtering', 'Qdrant Cloud'],
    ['Time-Series DB', 'InfluxDB OSS v2', 'Purpose-built, downsampling, continuous queries', 'InfluxDB Cloud'],
    ['Cache', 'Redis', 'Sub-ms reads, Pub/Sub, TTL', 'ElastiCache'],
    ['Object Storage', 'MinIO / S3', 'S3-compatible', 'Self-hosted -> S3'],
    ['LLM Serving', 'Ollama -> Together AI', 'Local for dev, cloud for prod', 'Together AI / Bedrock'],
    ['CI/CD', 'GitHub Actions', 'Free for public repos, matrix builds', 'GitHub'],
    ['Infra as Code', 'Terraform (HCL)', 'Reproducible, reviewable', 'Terraform Cloud'],
    ['Monitoring', 'Grafana + Prometheus + Loki', 'Metrics + logs + traces', 'Grafana Cloud'],
]
add_table(['Category', 'Technology', 'Justification', 'Production Service'], stack)

# ═══════════════════════════════════════════════════════════════════════════
# 20. AI MODELS
# ═══════════════════════════════════════════════════════════════════════════
add_heading('20. AI Models', level=1)

models = [
    ['Llama 3 8B', 'LLM', 'Meta', '8B parameters, 8K context, open-weight', 'SHELDRA reasoning, RAG generation, explanation', 'Fine-tunable on safety data; 4-bit quantized for 8GB VRAM'],
    ['E5-mistral-7b', 'Embedding', 'Microsoft', '4096-dim, Mistral-7b backbone', 'Document embedding for RAG', 'FP16 quantized; 50ms per embedding on CPU'],
    ['YOLOv8n', 'Detection', 'Ultralytics', '3.2M params, 30+ FPS on CPU', 'PPE detection (6 classes)', 'ONNX export; int8 quantized for edge'],
    ['YOLOv8s', 'Detection', 'Ultralytics', '11.2M params', 'Person detection + tracking', 'ONNX export; GPU optional'],
    ['RTMPose-m', 'Pose', 'OpenMMLab', '~5ms/person on GPU', 'Behavior analysis', 'ONNX export; 17-keypoint'],
    ['TimesNet', 'Time-Series', 'Microsoft Research', 'ICLR 2023 Best Paper', 'Anomaly detection', 'ONNX export; 60-input window'],
    ['Whisper tiny', 'STT', 'OpenAI', '39M params, multilingual', 'Voice input', 'ONNX export; browser or server'],
    ['Coqui TTS', 'TTS', 'Coqui AI', 'MIT license, offline', 'Voice output', 'Local inference; 2x real-time'],
]
add_table(['Model', 'Type', 'Source', 'Specs', 'Use', 'Deployment'], models)

# ═══════════════════════════════════════════════════════════════════════════
# 21. PERFORMANCE
# ═══════════════════════════════════════════════════════════════════════════
add_heading('21. Performance', level=1)

perf = [
    ['SHELDRA end-to-end (text chat)', '< 3s p95', 'LLM inference dominates (~2s). RAG retrieval < 200ms. KG query < 100ms. Streaming first token < 500ms (SSE).'],
    ['SHELDRA end-to-end (voice)', '< 5s p95', 'STT (~1s) + LLM (~2s) + TTS (~1s). Optimize with streaming STT (incremental) and pre-fetch TTS on partial response.'],
    ['Vision inference', '< 30ms per frame', 'YOLOv8n ONNX CPU. 5 FPS = 200ms per frame budget. Parallel inference across cameras (thread pool).'],
    ['Time-series anomaly', '< 100ms per window', 'TimesNet ONNX CPU. 60-point window. 1-second cadence = 900ms budget.'],
    ['Decision Tree accuracy', '< 2m per update cycle', 'Score aggregation + branch analysis. 5-minute cycle budget is generous. Tree size < 100 nodes.'],
    ['API throughput', '100+ req/s per instance', 'FastAPI async handlers. Stateless servers. Horizontal scaling via ECS.'],
    ['WebSocket fan-out', '10,000+ concurrent', 'Single WebSocket server (FastAPI) with async event loop. Scale via Redis Pub/Sub + multiple WS server instances.'],
    ['Database queries', '< 50ms p99 (KG)', 'Neo4j query cache (Redis). Frequent queries < 5ms. Complex multi-hop queries < 50ms.'],
]
add_table(['Metric', 'Target', 'Architecture Contribution'], perf)

# ═══════════════════════════════════════════════════════════════════════════
# 22. SCALABILITY
# ═══════════════════════════════════════════════════════════════════════════
add_heading('22. Scalability', level=1)

scalability = [
    ['Single Facility', '10-50 cameras, 200-1000 sensors, 50-200 workers',
     'Single SHELDRA Engine instance (4 vCPU, 16GB). Vision on CPU (ONNX). '
     'All databases on single nodes. Kafka with 1 partition per camera group.'],
    ['Multi-Facility (Mid)', '5-20 facilities, 500-5000 sensors, 500-2000 workers',
     'One SHELDRA Engine per facility (federated). Facility-level KG with cross-facility '
     'aggregation for benchmarking. Decision trees shared across facilities with same '
     'equipment/hazard profiles. Centralized RAG with per-facility document collections.'],
    ['Enterprise (100+ facilities)', '100-1000 facilities, 50K+ sensors, 20K+ workers',
     'Hierarchical SHELDRA: facility-level engines report to regional aggregators. '
     'Federated KG with global schema + local extensions. Global Decision Tree optimizer '
     'shares high-accuracy branches across facilities. Tiered storage: hot (InfluxDB), '
     'warm (downsampled), cold (S3 Parquet). GPU inference pool shared across facilities.'],
]
add_table(['Scale', 'Characteristics', 'Architecture'], scalability)

add_para('')
add_para(
    'The primary scalability bottleneck is the LLM Reasoning Core, which must process '
    'sequentially for each request. Mitigation: (1) request batching when multiple events '
    'occur simultaneously, (2) priority queuing (emergency events skip the queue), '
    '(3) smaller distilled models (Llama 3.2 3B) for low-complexity interactions with '
    'fallback to 8B for complex reasoning.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 23. DEMO FLOW
# ═══════════════════════════════════════════════════════════════════════════
add_heading('23. Demo Flow', level=1)

add_para('The 5-minute demonstration is structured in 8 contiguous segments:', bold=True)
demo_flow = [
    ['1. SHELDRA Introduction', '0:00-0:40', 'Dashboard loads. SHELDRA avatar appears with idle animation. SHELDRA speaks: "Welcome to Shift 3. I am SHELDRA, your safety coach. All systems normal." Safety score: 92. Normal factory operation visible in Digital Twin.'],
    ['2. Vision-Triggered Correction', '0:40-1:30', 'Synthetic camera feed shows worker entering without safety glasses. Vision Intelligence detects: {class: "missing_glasses", worker_id: "w3124", confidence: 0.94}. SHELDRA speaks: "Worker #3124, your safety glasses are on your helmet. Please adjust them before entering Zone A." Decision Tree flow chart opens beside avatar, highlighting the PPE correction path.'],
    ['3. Personalization Comparison', '1:30-2:15', 'Dashboard switches to novice worker profile. Same PPE violation triggers longer response with rationale. Novice: "Safety glasses protect your eyes from flying debris. Please wear them at all times in Zone A." Switch to expert: brief confirmation only. "One SHELDRA, three different responses."'],
    ['4. Sensor-Driven Emergency', '2:15-3:00', 'IoT simulator injects gas leak anomaly. Monitoring Intelligence detects: compound anomaly score 0.88. Decision Tree loads gas leak response tree. SHELDRA switches to emergency tone: "ATTENTION: Gas leak detected in Zone B. Evacuate via Exit 4 immediately." Tree path highlights in real-time. After resolution, tree updates: "Tree accuracy: 92%. Version 8 deployed."'],
    ['5. Accuracy Feedback Loop', '3:00-3:30', 'Accuracy dashboard: tree accuracy trend (86% -> 92% over 7 versions), branch breakdown (best/worst performing), version history. "This tree learns every 5 minutes. Every incident makes it smarter."'],
    ['6. VR Training', '3:30-4:00', 'Switch to VR scene. SHELDRA avatar in VR. Confined space training scenario. SHELDRA guides step-by-step. Worker response scored. Score updates profile. "Same SHELDRA that guides on the floor also trains in VR."'],
    ['7. Architecture Reveal', '4:00-4:30', 'System architecture diagram displayed. "Everything you saw \u2014 one Intelligence Engine, 11 modules. The avatar is just an interface. SHELDRA is the brain."'],
    ['8. Closing & Q&A Hook', '4:30-5:00', '"SHELDRA is the first AI safety coach that walks beside every worker. Built in 72 hours. Designed for any facility. We\u2019d love to discuss how this redefines industrial safety."'],
]
add_table(['Segment', 'Duration', 'Content'], demo_flow)

# ═══════════════════════════════════════════════════════════════════════════
# 24. ROADMAP
# ═══════════════════════════════════════════════════════════════════════════
add_heading('24. Roadmap', level=1)

roadmap = [
    ['MVP (Hackathon)', 'SHELDRA Intelligence Engine with 11 modules, 3 worker profiles, 5 Decision Trees, synthetic data, single-facility, web dashboard'],
    ['Alpha (Month 1-2)', 'Real camera + IoT integration, auto-generated Decision Trees, learning speed detection, VR training, real-data validation'],
    ['Beta (Month 3-4)', 'Multi-facility, federated KGs, mobile app, AR overlay, SSO, SOC 2 prep, 3+ pilot customers'],
    ['GA v1 (Month 5-7)', 'On-premise deployment, air-gapped mode, custom model fine-tuning, DT marketplace, 99.9% SLA, 10+ customers'],
    ['Enterprise SaaS (Month 8-12)', 'Multi-region, cross-facility benchmarking, insurance API, regulatory filing, white-label, 50+ customers'],
    ['Platform (Year 2+)', 'Industry-specific models, autonomous actions, regulator-approved safety twin, community ecosystem, $50M+ ARR'],
]
add_table(['Phase', 'Capabilities'], roadmap)

# ═══════════════════════════════════════════════════════════════════════════
# 25. RISKS
# ═══════════════════════════════════════════════════════════════════════════
add_heading('25. Risks', level=1)

risks = [
    ['LLM Hallucination', 'Critical', 'High', 'SHELDRA generates incorrect safety procedure.', 'RAG grounding with citation anchoring. Confidence threshold with escalation fallback. Human review of all procedural guidance. Multiple guardrail layers.'],
    ['Vision Accuracy in Edge Cases', 'High', 'Medium', 'PPE detection fails in low light, occlusion, or unusual angles.', 'Confidence filtering (min 0.5). Human-in-the-loop for low-confidence detections. Ensemble of YOLOv8 + rule-based checks.'],
    ['Integration Complexity', 'High', 'Medium', 'Each facility has different cameras, PLCs, and sensor protocols.', 'Adapter pattern from day one. REST + MQTT bridges. Start with synthetic data; real integration in alpha phase.'],
    ['3D Avatar Performance', 'Medium', 'Medium', 'Browser-based 3D avatar lags on low-end hardware.', '2D Lottie fallback with identical behavior. Progressive enhancement: 3D if GPU available, 2D otherwise.'],
    ['Data Privacy Regulations', 'Medium', 'High', 'Camera feeds contain biometric data. Sensor data reveals operational patterns.', 'On-premise deployment option. Edge processing (no raw video leaves facility). Data retention policies. Compliance-by-design from day one.'],
    ['Market Adoption', 'Medium', 'Medium', 'Industrial companies are risk-averse and slow to adopt new technology.', 'Start with "safety co-pilot" (low-risk, high-value). SOC 2 from day one. Case studies with measurable ROI. Free tier for small facilities.'],
]
add_table(['Risk', 'Impact', 'Prob.', 'Description', 'Mitigation'], risks)

# ═══════════════════════════════════════════════════════════════════════════
# 26. FUTURE ENHANCEMENTS
# ═══════════════════════════════════════════════════════════════════════════
add_heading('26. Future Enhancements', level=1)

enhancements = [
    ['Acoustic Intelligence', 'Microphone array module for abnormal sound detection (gas leak hiss, equipment bearing failure whine, human distress calls). Reuses existing EventOrchestrator pattern.'],
    ['Wearable Integration', 'Biometric sensors (heart rate, skin temperature, galvanic skin response) for stress and fatigue detection. SHELDRA adapts response based on physiological state.'],
    ['Predictive Maintenance', 'Extend TimesNet to predict equipment failure before it occurs (currently detects anomalies after they begin). Integrate with Decision Tree for pre-failure procedural guidance.'],
    ['Cross-Facility Learning', 'Privacy-preserving federated learning across facilities. Decision Tree improvements at one facility propagate to others without sharing raw incident data.'],
    ['Autonomous Safety Actions', 'For extreme scenarios (imminent catastrophic failure), SHELDRA can initiate safety actions (equipment shutdown, zone evacuation) without waiting for supervisor confirmation. Configurable per facility risk tolerance.'],
    ['Regulatory Filing Automation', 'Auto-generate OSHA 300 logs, incident reports, and compliance documentation from SHELDRA interaction history. Reduce administrative burden on safety officers.'],
    ['Multi-Language Real-Time Translation', 'SHELDRA detects worker language and responds in kind. Enables multilingual teams to receive guidance in their preferred language.'],
]
for name, desc in enhancements:
    p = doc.add_paragraph()
    r = p.add_run(name + ': ')
    r.bold = True
    p.add_run(desc)

# ── Save ──
output_path = '/Users/nayantraramakrishnan/Desktop/developer/Projects/ey-hackathon/industrial-safety-intelligence-platform-architecture.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
